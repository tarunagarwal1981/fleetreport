import streamlit as st
import requests
import pandas as pd
import json
import io
from datetime import datetime, timedelta
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import PatternFill, Font, Alignment, Color
import time
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from collections import defaultdict

# Page configuration with improved styling
st.set_page_config(
    page_title="Vessel Performance Report Tool",
    page_icon="🚢",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for better UI
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(120deg, #16222A 0%, #3A6073 100%);
        padding: 2.2rem;
        border-radius: 18px;
        margin-bottom: 1.5rem;
        text-align: center;
        color: white;
        box-shadow: 0 15px 35px rgba(0,0,0,0.15);
    }
    .section-header {
        background: linear-gradient(90deg, #3A6073 0%, #16222A 100%);
        padding: 0.9rem 1.2rem;
        border-radius: 12px;
        margin: 1.5rem 0 0.8rem 0;
        color: white;
        font-weight: 600;
        letter-spacing: 0.5px;
    }
    .summary-bar {
        background: #ffffff;
        border: 1px solid rgba(22,34,42,0.08);
        border-radius: 16px;
        padding: 1rem;
        display: flex;
        gap: 1rem;
        flex-wrap: wrap;
        box-shadow: 0 10px 25px rgba(0,0,0,0.08);
        margin-bottom: 1.2rem;
    }
    .summary-chip {
        flex: 1 1 180px;
        min-width: 140px;
        background: #f6f8fb;
        border-radius: 12px;
        padding: 0.75rem 1rem;
        border: 1px solid rgba(58,96,115,0.12);
    }
    .summary-chip h4 {
        margin: 0;
        font-size: 0.9rem;
        color: #5f6c7b;
        font-weight: 500;
    }
    .summary-chip p {
        margin: 0.2rem 0 0 0;
        font-size: 1.35rem;
        font-weight: 600;
        color: #16222A;
    }
    .metric-card {
        background: #ffffff;
        padding: 1rem;
        border-radius: 12px;
        border: 1px solid rgba(58,96,115,0.12);
        box-shadow: 0 8px 18px rgba(0,0,0,0.06);
        margin: 0.4rem 0;
    }
    .doc-pill {
        display: inline-flex;
        align-items: center;
        gap: 0.4rem;
        padding: 0.35rem 0.9rem;
        background: rgba(58,96,115,0.12);
        border-radius: 999px;
        font-size: 0.85rem;
        color: #16222A;
        font-weight: 600;
    }
    .stButton > button {
        background: linear-gradient(120deg, #0f2027, #203a43, #2c5364);
        color: white;
        border: none;
        border-radius: 10px;
        padding: 0.55rem 1.25rem;
        font-weight: 600;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 12px 20px rgba(32,58,67,0.25);
    }
    .success-box, .info-box, .alert-box {
        border-radius: 12px;
        padding: 0.9rem 1.1rem;
        margin: 0.8rem 0;
        color: #ffffff;
    }
    .success-box { background: linear-gradient(120deg, #11998e 0%, #38ef7d 100%); }
    .info-box { background: linear-gradient(120deg, #396afc 0%, #2948ff 100%); }
    .alert-box { background: linear-gradient(120deg, #ff512f 0%, #dd2476 100%); }
    @media (max-width: 768px) {
        .main-header { padding: 1.5rem; }
        .summary-bar { flex-direction: column; }
        .summary-chip { width: 100%; }
        .stTabs [role="tablist"] > div { font-size: 0.85rem; }
    }
</style>
""", unsafe_allow_html=True)

MAX_DIAGNOSTIC_ENTRIES = 8


def log_diagnostic_event(label, status, detail):
    timestamp = datetime.now().strftime("%H:%M:%S")
    entry = {
        "label": label,
        "status": status,
        "detail": detail,
        "time": timestamp,
    }
    st.session_state.diagnostics.insert(0, entry)
    st.session_state.diagnostics = st.session_state.diagnostics[:MAX_DIAGNOSTIC_ENTRIES]


def render_summary_bar(container):
    metrics = st.session_state.summary_metrics
    container.markdown(f"""
        <div class="summary-bar">
            <div class="summary-chip">
                <h4>DOC</h4>
                <p>{metrics.get('doc', 'All Offices')}</p>
            </div>
            <div class="summary-chip">
                <h4>Total Vessels</h4>
                <p>{metrics.get('total', 0)}</p>
            </div>
            <div class="summary-chip">
                <h4>Filtered</h4>
                <p>{metrics.get('filtered', 0)}</p>
            </div>
            <div class="summary-chip">
                <h4>Selected</h4>
                <p>{metrics.get('selected', 0)}</p>
            </div>
        </div>
""", unsafe_allow_html=True)

# Initialize session state
if 'vessels' not in st.session_state:
    st.session_state.vessels = []

if 'selected_vessels' not in st.session_state or not isinstance(st.session_state.selected_vessels, set):
    st.session_state.selected_vessels = set()

if 'report_data' not in st.session_state:
    st.session_state.report_data = None

if 'search_query' not in st.session_state:
    st.session_state.search_query = ""

if 'report_months' not in st.session_state:
    st.session_state.report_months = 2

if 'selected_office' not in st.session_state:
    st.session_state.selected_office = "All Offices"

if 'diagnostics' not in st.session_state:
    st.session_state.diagnostics = []

if 'summary_metrics' not in st.session_state:
    st.session_state.summary_metrics = {
        "doc": "All Offices",
        "total": 0,
        "filtered": 0,
        "selected": 0,
    }

if 'checkbox_version' not in st.session_state:
    st.session_state.checkbox_version = 0

# Enhanced Lambda Invocation Helper
def invoke_lambda_function_url(lambda_url, payload, timeout=60):
    """Invoke Lambda function via its Function URL using HTTP POST with performance tracking."""
    try:
        start_time = time.time()
        headers = {'Content-Type': 'application/json'}
        json_payload = json.dumps(payload)

        response = requests.post(
            lambda_url,
            headers=headers,
            data=json_payload,
            timeout=timeout
        )

        response_time = time.time() - start_time

        if response.status_code != 200:
            st.error(f"HTTP error: {response.status_code} {response.reason} for url: {lambda_url}")
            log_diagnostic_event("HTTP Error", "error", f"{response.status_code} {response.reason}")
            return None

        result = response.json()
        st.success(f"✅ Data retrieved in {response_time:.2f}s")
        log_diagnostic_event("Lambda Query", "success", f"{response_time:.2f}s")
        return result

    except requests.exceptions.HTTPError as http_err:
        st.error(f"HTTP error: {http_err}")
        log_diagnostic_event("HTTP Error", "error", str(http_err))
        return None
    except requests.exceptions.ConnectionError as conn_err:
        st.error(f"Connection error: {conn_err}")
        log_diagnostic_event("Connection Error", "error", str(conn_err))
        return None
    except requests.exceptions.Timeout as timeout_err:
        st.error(f"Timeout error: {timeout_err}")
        log_diagnostic_event("Timeout", "error", str(timeout_err))
        return None
    except requests.exceptions.RequestException as req_err:
        st.error(f"Request error: {req_err}")
        log_diagnostic_event("Request Error", "error", str(req_err))
        return None
    except Exception as e:
        st.error(f"Unexpected error: {str(e)}")
        log_diagnostic_event("Unexpected Error", "error", str(e))
        return None

# Cached vessel loading function
@st.cache_data(ttl=3600)
def fetch_vessel_directory(lambda_url):
    """Fetch vessel names with office information."""
    office_query = """
        SELECT 
            vessel_name, 
            COALESCE(office_doc, 'Unassigned DOC') AS office
        FROM vessel_particulars
        WHERE vessel_name IS NOT NULL
        ORDER BY office, vessel_name
        LIMIT 1200
    """

    result = invoke_lambda_function_url(lambda_url, {"sql_query": office_query})

    # Fallback: if office column is unavailable, retry with legacy query
    if not result:
        return []
    if isinstance(result, dict) and result.get("error"):
        return fallback_vessel_directory(lambda_url)
    if isinstance(result, list) and result and isinstance(result[0], dict) and "office" not in result[0]:
        return fallback_vessel_directory(lambda_url)

    cleaned_records = []
    for item in result:
        if isinstance(item, dict):
            vessel = item.get('vessel_name')
            office = item.get('office') or "Unassigned DOC"
        else:
            vessel = item[0] if isinstance(item, (list, tuple)) and item else None
            office = item[1] if isinstance(item, (list, tuple)) and len(item) > 1 else "Unassigned DOC"
        if vessel:
            cleaned_records.append({"vessel_name": vessel, "office": office})

    return cleaned_records


def fallback_vessel_directory(lambda_url):
    """Legacy loader when office column is unavailable."""
    legacy_query = "SELECT vessel_name FROM vessel_particulars ORDER BY vessel_name LIMIT 1200"
    result = invoke_lambda_function_url(lambda_url, {"sql_query": legacy_query})
    if not result:
        st.warning("DOC column not available; falling back to basic vessel list.")
        return []
    cleaned_records = []
    for item in result:
        vessel = None
        if isinstance(item, dict) and 'vessel_name' in item:
            vessel = item['vessel_name']
        elif isinstance(item, str):
            vessel = item
        elif isinstance(item, (list, tuple)) and item:
            vessel = item[0]
        if vessel:
            cleaned_records.append({"vessel_name": vessel, "office": "Unassigned DOC"})
    return cleaned_records

def filter_vessels_client_side(vessels, search_term):
    """Filter vessels on client side for better responsiveness."""
    if not search_term:
        return vessels
  
    search_lower = search_term.lower()
    return [v for v in vessels if search_lower in v.lower()]

def query_report_data(lambda_url, vessel_names, num_months):
    """Enhanced version of the original query_report_data function with better progress tracking and fixed date logic."""
    if not vessel_names:
        return pd.DataFrame()

    today = datetime.now()
    # Get the first day of the current month
    first_day_current_month = today.replace(day=1)

    # Prepare date strings and column names based on num_months
    hull_dates_info = []
    me_dates_info = []

    for i in range(num_months):
        # Calculate the target month (going backwards from current month)
        # For i=0: current month - 1 (previous month)
        # For i=1: current month - 2 (two months ago)
        months_back = i + 1
        
        # Get the first day of the target month
        if first_day_current_month.month <= months_back:
            # Need to go back to previous year
            target_year = first_day_current_month.year - 1
            target_month = 12 - (months_back - first_day_current_month.month)
        else:
            target_year = first_day_current_month.year
            target_month = first_day_current_month.month - months_back
        
        target_month_first_day = datetime(target_year, target_month, 1)
        
        # Hull Condition Dates (last day of the target month)
        if target_month == 12:
            next_month_first = datetime(target_year + 1, 1, 1)
        else:
            next_month_first = datetime(target_year, target_month + 1, 1)
        
        target_month_last_day = next_month_first - timedelta(days=1)

        hull_date_str_start = target_month_first_day.strftime("%Y-%m-%d")
        hull_date_str_end = target_month_last_day.strftime("%Y-%m-%d")
        hull_col_name = f"Hull Condition {target_month_last_day.strftime('%b %y')}"
        hull_power_loss_col_name = f"Excess Power % {target_month_last_day.strftime('%b %y')}"
        hull_dates_info.append({
            'date_str_start': hull_date_str_start,
            'date_str_end': hull_date_str_end,
            'col_name': hull_col_name,
            'power_loss_col_name': hull_power_loss_col_name,
            'months_back': months_back
        })

        # ME SFOC Dates (average of the entire target month)
        me_col_name = f"ME Efficiency {target_month_last_day.strftime('%b %y')}"
        me_dates_info.append({
            'col_name': me_col_name,
            'months_back': months_back,
            'target_month_first': target_month_first_day
        })

    # Process vessels in smaller batches with enhanced progress tracking
    batch_size = 10
    all_fuel_saving_data = []
    all_cii_data = []
    all_hull_data_by_month = {info['power_loss_col_name']: [] for info in hull_dates_info}
    all_me_data_by_month = {info['col_name']: [] for info in me_dates_info}

    total_batches = (len(vessel_names) + batch_size - 1) // batch_size

    # Create progress bar and status text
    progress_bar = st.progress(0)
    status_text = st.empty()

    for i in range(0, len(vessel_names), batch_size):
        batch_vessels = vessel_names[i:i+batch_size]
        batch_num = i//batch_size + 1

        # Update progress
        progress = batch_num / total_batches
        progress_bar.progress(progress)
        status_text.info(f"🔄 Processing batch {batch_num} of {total_batches} ({len(batch_vessels)} vessels)")

        quoted_vessel_names = [f"'{name}'" for name in batch_vessels]
        vessel_names_list_str = ", ".join(quoted_vessel_names)

        batch_queries = []

        # Hull Roughness queries - get most recent data within the target month
        for hull_info in hull_dates_info:
            batch_queries.append((hull_info['power_loss_col_name'], f"""
SELECT vessel_name, hull_rough_power_loss_pct_ed
FROM (
    SELECT vessel_name, hull_rough_power_loss_pct_ed,
           ROW_NUMBER() OVER (PARTITION BY vessel_name ORDER BY updated_ts DESC) as rn
    FROM hull_performance_six_months_daily
    WHERE vessel_name IN ({vessel_names_list_str})
    AND CAST(updated_ts AS DATE) >= '{hull_info['date_str_start']}'
    AND CAST(updated_ts AS DATE) <= '{hull_info['date_str_end']}'
) AS subquery
WHERE rn = 1
""", all_hull_data_by_month[hull_info['power_loss_col_name']]))

        # ME SFOC queries with corrected date logic
        for me_info in me_dates_info:
            batch_queries.append((me_info['col_name'], f"""
SELECT vp.vessel_name, AVG(vps.me_sfoc) AS avg_me_sfoc
FROM vessel_performance_summary vps
JOIN vessel_particulars vp ON CAST(vps.vessel_imo AS TEXT) = CAST(vp.vessel_imo AS TEXT)
WHERE vp.vessel_name IN ({vessel_names_list_str})
AND vps.reportdate >= DATE_TRUNC('month', CURRENT_DATE - INTERVAL '{me_info['months_back']} month')
AND vps.reportdate < DATE_TRUNC('month', CURRENT_DATE - INTERVAL '{me_info['months_back'] - 1} month')
GROUP BY vp.vessel_name
""", all_me_data_by_month[me_info['col_name']]))

        # Fixed queries
        batch_queries.append(("Potential Fuel Saving", f"""
SELECT vessel_name, hull_rough_excess_consumption_mt_ed
FROM hull_performance_six_months
WHERE vessel_name IN ({vessel_names_list_str})
""", all_fuel_saving_data))

        batch_queries.append(("YTD CII", f"""
SELECT vp.vessel_name, cy.cii_rating
FROM vessel_particulars vp
JOIN cii_ytd cy ON CAST(vp.vessel_imo AS TEXT) = CAST(cy.vessel_imo AS TEXT)
WHERE vp.vessel_name IN ({vessel_names_list_str})
""", all_cii_data))

        # Execute each query
        for query_name, query, data_list in batch_queries:
            with st.spinner(f"Fetching {query_name} data..."):
                result = invoke_lambda_function_url(lambda_url, {"sql_query": query})
                if result:
                    data_list.extend(result)

    # Clear progress indicators
    progress_bar.empty()
    status_text.empty()

    # Process all collected data
    df_final = pd.DataFrame({'Vessel Name': list(vessel_names)})

    # Hull Data processing and merging
    for hull_info in hull_dates_info:
        df_hull = pd.DataFrame()
        if all_hull_data_by_month[hull_info['power_loss_col_name']]:
            try:
                df_hull = pd.DataFrame(all_hull_data_by_month[hull_info['power_loss_col_name']])
                if 'hull_rough_power_loss_pct_ed' in df_hull.columns:
                    df_hull = df_hull.rename(columns={'hull_rough_power_loss_pct_ed': hull_info['power_loss_col_name']})
                else:
                    df_hull[hull_info['power_loss_col_name']] = pd.NA
                df_hull = df_hull.rename(columns={'vessel_name': 'Vessel Name'})
            except Exception as e:
                st.error(f"Error processing {hull_info['col_name']} data: {str(e)}")
                df_hull = pd.DataFrame()
        if not df_hull.empty:
            df_final = pd.merge(df_final, df_hull, on='Vessel Name', how='left')

    # ME Data processing and merging
    for me_info in me_dates_info:
        df_me = pd.DataFrame()
        if all_me_data_by_month[me_info['col_name']]:
            try:
                df_me = pd.DataFrame(all_me_data_by_month[me_info['col_name']])
                if 'avg_me_sfoc' in df_me.columns:
                    df_me = df_me.rename(columns={'avg_me_sfoc': me_info['col_name']})
                else:
                    df_me[me_info['col_name']] = pd.NA
                df_me = df_me.rename(columns={'vessel_name': 'Vessel Name'})
            except Exception as e:
                st.error(f"Error processing {me_info['col_name']} data: {str(e)}")
                df_me = pd.DataFrame()
        if not df_me.empty:
            df_final = pd.merge(df_final, df_me, on='Vessel Name', how='left')

    # Fuel saving and CII data processing
    df_fuel_saving = pd.DataFrame()
    if all_fuel_saving_data:
        try:
            df_fuel_saving = pd.DataFrame(all_fuel_saving_data)
            if 'hull_rough_excess_consumption_mt_ed' in df_fuel_saving.columns:
                df_fuel_saving = df_fuel_saving.rename(columns={'hull_rough_excess_consumption_mt_ed': 'Potential Fuel Saving (MT/Day)'})
                df_fuel_saving['Potential Fuel Saving (MT/Day)'] = df_fuel_saving['Potential Fuel Saving (MT/Day)'].apply(
                    lambda x: 4.9 if pd.notna(x) and x > 5 else (0.0 if pd.notna(x) and x < 0 else x)
                ).round(2)
            else:
                df_fuel_saving['Potential Fuel Saving (MT/Day)'] = pd.NA
            df_fuel_saving = df_fuel_saving.rename(columns={'vessel_name': 'Vessel Name'})
        except Exception as e:
            st.error(f"Error processing fuel saving data: {str(e)}")
            df_fuel_saving = pd.DataFrame()

    df_cii = pd.DataFrame()
    if all_cii_data:
        try:
            df_cii = pd.DataFrame(all_cii_data)
            if 'cii_rating' in df_cii.columns:
                df_cii = df_cii.rename(columns={'cii_rating': 'YTD CII'})
            else:
                df_cii['YTD CII'] = pd.NA
            df_cii = df_cii.rename(columns={'vessel_name': 'Vessel Name'})
        except Exception as e:
            st.error(f"Error processing CII data: {str(e)}")
            df_cii = pd.DataFrame()

    # Merge other data
    if not df_fuel_saving.empty:
        df_final = pd.merge(df_final, df_fuel_saving, on='Vessel Name', how='left')

    if not df_cii.empty:
        df_final = pd.merge(df_final, df_cii, on='Vessel Name', how='left')

    if df_final.empty:
        return pd.DataFrame()

    # Post-merge processing for final report
    df_final.insert(0, 'S. No.', range(1, 1 + len(df_final)))

    # Hull Condition logic
    def get_hull_condition(value):
        if pd.isna(value):
            return "N/A"
        if value < 15:
            return "Good"
        elif 15 <= value <= 25:
            return "Average"
        else:
            return "Poor"

    # Apply Hull Condition to historical columns and format power loss values
    for hull_info in hull_dates_info:
        if hull_info['power_loss_col_name'] in df_final.columns:
            # Create Hull Condition column (Good/Average/Poor)
            df_final[hull_info['col_name']] = df_final[hull_info['power_loss_col_name']].apply(get_hull_condition)
            # Format power loss percentage column (round to 1 decimal, keep as number for sorting)
            df_final[hull_info['power_loss_col_name']] = df_final[hull_info['power_loss_col_name']].apply(
                lambda x: round(x, 1) if pd.notna(x) else pd.NA
            )
        else:
            df_final[hull_info['col_name']] = "N/A"
            df_final[hull_info['power_loss_col_name']] = pd.NA

    # ME Efficiency logic and comments
    def get_me_efficiency_and_comment(value):
        if pd.isna(value):
            return "N/A", ""
        if value < 160:
            return "Anomalous data", f"SFOC value is {value:.1f} g/kWh, unusually low, indicating potential data anomaly."
        elif value < 180:
            return "Good", ""
        elif 180 <= value <= 190:
            return "Average", ""
        else:
            return "Poor", ""

    # Initialize a temporary column for comments
    df_final['temp_comments_list'] = [[] for _ in range(len(df_final))]

    # Apply ME Efficiency and populate comments
    for me_info in me_dates_info:
        if me_info['col_name'] in df_final.columns:
            # Apply the function to get both status and comment
            df_final[[me_info['col_name'], 'current_me_comment']] = df_final[me_info['col_name']].apply(
                lambda x: pd.Series(get_me_efficiency_and_comment(x))
            )
            # Append comments for anomalous data to the temporary list
            df_final['temp_comments_list'] = df_final.apply(
                lambda row: row['temp_comments_list'] + [f"ME Efficiency ({me_info['col_name'].split(' ')[-2:]}): {row['current_me_comment']}"] if row['current_me_comment'] else row['temp_comments_list'],
                axis=1
            )
            df_final = df_final.drop(columns=['current_me_comment'])

    # Join all collected comments into the final 'Comments' column
    df_final['Comments'] = df_final['temp_comments_list'].apply(lambda x: " ".join(x).strip())
    df_final = df_final.drop(columns=['temp_comments_list'])

    # Define the desired order of columns (hide Potential Fuel Saving)
    # Order: S. No., Vessel Name, All Power Loss % columns, All Hull Condition columns, ME Efficiency, YTD CII, Comments
    desired_columns_order = ['S. No.', 'Vessel Name']
    # Add all Power Loss % columns first (for each month)
    for hull_info in hull_dates_info:
        desired_columns_order.append(hull_info['power_loss_col_name'])  # Power Loss % (numeric value)
    # Then add all Hull Condition columns (for each month)
    for hull_info in hull_dates_info:
        desired_columns_order.append(hull_info['col_name'])  # Hull Condition (Good/Average/Poor)
    # Add ME Efficiency columns
    for me_info in me_dates_info:
        desired_columns_order.append(me_info['col_name'])
    # desired_columns_order.append('Potential Fuel Saving (MT/Day)')  # Hidden
    desired_columns_order.extend(['YTD CII', 'Comments'])

    # Filter df_final to only include columns that exist and are in the desired order
    existing_and_ordered_columns = [col for col in desired_columns_order if col in df_final.columns]
    df_final = df_final[existing_and_ordered_columns]

    st.success("✅ Enhanced report data retrieved and processed successfully!")
    return df_final

# Enhanced styling function with CII color coding
def style_condition_columns(row):
    """Apply styling to condition columns including CII rating text color."""
    styles = [''] * len(row)

    # Style hull condition columns
    hull_condition_cols = [col for col in row.index if 'Hull Condition' in col and 'Power Loss' not in col]
    for col_name in hull_condition_cols:
        if col_name in row.index:
            hull_val = row[col_name]
            if hull_val == "Good":
                styles[row.index.get_loc(col_name)] = 'background-color: #d4edda; color: black;'
            elif hull_val == "Average":
                styles[row.index.get_loc(col_name)] = 'background-color: #fff3cd; color: black;'
            elif hull_val == "Poor":
                styles[row.index.get_loc(col_name)] = 'background-color: #f8d7da; color: black;'
    
    # Style hull power loss percentage columns (based on numeric values)
    hull_power_loss_cols = [col for col in row.index if 'Excess Power %' in col]
    for col_name in hull_power_loss_cols:
        if col_name in row.index:
            power_loss_val = row[col_name]
            if pd.notna(power_loss_val):
                try:
                    power_loss_num = float(power_loss_val)
                    if power_loss_num < 15:
                        styles[row.index.get_loc(col_name)] = 'background-color: #d4edda; color: black; font-weight: 500;'
                    elif 15 <= power_loss_num <= 25:
                        styles[row.index.get_loc(col_name)] = 'background-color: #fff3cd; color: black; font-weight: 500;'
                    else:
                        styles[row.index.get_loc(col_name)] = 'background-color: #f8d7da; color: black; font-weight: 500;'
                except (ValueError, TypeError):
                    pass

    # Style ME efficiency columns
    me_efficiency_cols = [col for col in row.index if 'ME Efficiency' in col]
    for col_name in me_efficiency_cols:
        if col_name in row.index:
            me_val = row[col_name]
            if me_val == "Good":
                styles[row.index.get_loc(col_name)] = 'background-color: #d4edda; color: black;'
            elif me_val == "Average":
                styles[row.index.get_loc(col_name)] = 'background-color: #fff3cd; color: black;'
            elif me_val == "Poor":
                styles[row.index.get_loc(col_name)] = 'background-color: #f8d7da; color: black;'
            elif me_val == "Anomalous data":
                styles[row.index.get_loc(col_name)] = 'background-color: #FF0000; color: white;'  # Red background

    # Style YTD CII column with background colors and black text
    if 'YTD CII' in row.index:
        cii_val = str(row['YTD CII']).upper() if pd.notna(row['YTD CII']) else "N/A"
        if cii_val == "A":
            styles[row.index.get_loc('YTD CII')] = 'background-color: #006400; color: black; font-weight: bold;'  # Dark green background
        elif cii_val == "B":
            styles[row.index.get_loc('YTD CII')] = 'background-color: #90EE90; color: black; font-weight: bold;'  # Light green background
        elif cii_val == "C":
            styles[row.index.get_loc('YTD CII')] = 'background-color: #FFD700; color: black; font-weight: bold;'  # Yellow background
        elif cii_val == "D":
            styles[row.index.get_loc('YTD CII')] = 'background-color: #FF8C00; color: black; font-weight: bold;'  # Orange background
        elif cii_val == "E":
            styles[row.index.get_loc('YTD CII')] = 'background-color: #FF0000; color: black; font-weight: bold;'  # Red background

    return styles

def create_excel_download_with_styling(df, filename):
    """Create Excel file with styling including CII color coding."""
    output = io.BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = "Vessel Report"

    # Write headers
    for col_idx, col_name in enumerate(df.columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(wrap_text=True, horizontal='center', vertical='center')

    # Write data and apply styling
    for row_idx, row_data in df.iterrows():
        for col_idx, (col_name, cell_value) in enumerate(row_data.items(), 1):
            # Convert pandas NA to None for Excel compatibility
            excel_value = None if pd.isna(cell_value) else cell_value
            cell = ws.cell(row=row_idx + 2, column=col_idx, value=excel_value)
            cell.alignment = Alignment(wrap_text=True, vertical='top')

            if 'Hull Condition' in col_name and 'Power Loss' not in col_name:
                # Hull Condition text columns (Good/Average/Poor)
                if cell_value == "Good":
                    cell.fill = PatternFill(start_color="D4EDDA", end_color="D4EDDA", fill_type="solid")
                elif cell_value == "Average":
                    cell.fill = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid")
                elif cell_value == "Poor":
                    cell.fill = PatternFill(start_color="F8D7DA", end_color="F8D7DA", fill_type="solid")
                cell.alignment = Alignment(wrap_text=True, horizontal='center', vertical='center')
            elif 'Excess Power %' in col_name:
                # Power Loss percentage columns (numeric values)
                if pd.notna(cell_value):
                    try:
                        power_loss_num = float(cell_value)
                        cell.number_format = '0.0'  # Format as number with 1 decimal
                        if power_loss_num < 15:
                            cell.fill = PatternFill(start_color="D4EDDA", end_color="D4EDDA", fill_type="solid")
                        elif 15 <= power_loss_num <= 25:
                            cell.fill = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid")
                        else:
                            cell.fill = PatternFill(start_color="F8D7DA", end_color="F8D7DA", fill_type="solid")
                    except (ValueError, TypeError):
                        pass
                cell.alignment = Alignment(wrap_text=True, horizontal='center', vertical='center')
            elif 'ME Efficiency' in col_name:
                if cell_value == "Good":
                    cell.fill = PatternFill(start_color="D4EDDA", end_color="D4EDDA", fill_type="solid")
                elif cell_value == "Average":
                    cell.fill = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid")
                elif cell_value == "Poor":
                    cell.fill = PatternFill(start_color="F8D7DA", end_color="F8D7DA", fill_type="solid")
                elif cell_value == "Anomalous data":
                    cell.fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
                    cell.font = Font(color="FFFFFF")  # White text on red background
                cell.alignment = Alignment(wrap_text=True, horizontal='center', vertical='center')
            elif col_name == 'YTD CII':
                # Apply CII background color coding
                cii_val = str(cell_value).upper() if pd.notna(cell_value) else "N/A"
                cell.font = Font(bold=True, color="000000")  # Black text
                if cii_val == "A":
                    cell.fill = PatternFill(start_color="006400", end_color="006400", fill_type="solid")  # Dark green
                elif cii_val == "B":
                    cell.fill = PatternFill(start_color="90EE90", end_color="90EE90", fill_type="solid")  # Light green
                elif cii_val == "C":
                    cell.fill = PatternFill(start_color="FFD700", end_color="FFD700", fill_type="solid")  # Yellow
                elif cii_val == "D":
                    cell.fill = PatternFill(start_color="FF8C00", end_color="FF8C00", fill_type="solid")  # Orange
                elif cii_val == "E":
                    cell.fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")  # Red
                cell.alignment = Alignment(wrap_text=True, horizontal='center', vertical='center')
            elif col_name == 'Comments':
                cell.alignment = Alignment(wrap_text=True, horizontal='left', vertical='top')
            elif col_name == 'Potential Fuel Saving (MT/Day)':
                cell.alignment = Alignment(wrap_text=True, horizontal='center', vertical='center')

    # Auto-adjust column widths and set specific width for Comments and S. No.
    for col_idx, column in enumerate(df.columns, 1):
        column_letter = get_column_letter(col_idx)
        if column == 'Comments':
            ws.column_dimensions[column_letter].width = 40
        elif column == 'S. No.':
            ws.column_dimensions[column_letter].width = 8
        else:
            max_length = 0
            for cell in ws[column_letter]:
                try:
                    if cell.value is not None and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2) * 1.2
            ws.column_dimensions[column_letter].width = adjusted_width

    wb.save(output)
    return output.getvalue()

def get_cell_color(cell_value):
    """Get background color for table cell based on value."""
    color_map = {
        "Good": "D4EDDA",
        "Average": "FFF3CD",
        "Poor": "F8D7DA",
        "Anomalous data": "#f8d7da"
    }
    return color_map.get(cell_value, None)

def get_cii_background_color(cii_value):
    """Get background color hex for CII rating."""
    cii_val = str(cii_value).upper() if pd.notna(cii_value) else "N/A"
    color_map = {
        "A": "006400",      # Dark green
        "B": "90EE90",      # Light green
        "C": "FFD700",      # Yellow
        "D": "FF8C00",      # Orange
        "E": "FF0000"       # Red
    }
    return color_map.get(cii_val, None)

def set_cell_border(cell, **kwargs):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for border_name in ("top", "left", "bottom", "right"):
        if border_name in kwargs:
            border_element = OxmlElement(f"w:{border_name}")
            for attr, value in kwargs[border_name].items():
                border_element.set(qn(f"w:{attr}"), str(value))
            tcPr.append(border_element)

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell using direct XML manipulation."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)

def create_enhanced_word_report(df, template_path="Fleet Performance Template.docx", num_months=2):
    """Create an enhanced Word report with improved table formatting and CII color coding."""
    try:
        if not os.path.exists(template_path):
            st.error(f"Template file '{template_path}' not found in the repository root.")
            return None

        doc = Document(template_path)
      
        # Find placeholder and replace with report
        placeholder_found = False
      
        for paragraph in doc.paragraphs:
            if "{{Template}}" in paragraph.text:
                paragraph.clear()
                placeholder_found = True
              
                # Add report title with better styling
                title_paragraph = doc.add_paragraph()
                title_run = title_paragraph.add_run("Fleet Performance Analysis Report")
                title_run.font.size = Pt(24)
                title_run.font.bold = True
                title_run.font.color.rgb = RGBColor(47, 117, 181)  # Blue color
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
              
                # Add generation date
                date_paragraph = doc.add_paragraph()
                date_run = date_paragraph.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
                date_run.font.italic = True
                date_run.font.size = Pt(12)
                date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
              
                # Add some spacing
                doc.add_paragraph("")
              
                # Create table with simplified structure
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
              
                # Set table width to page width
                table.autofit = False
                table.allow_autofit = False
              
                # Define column widths based on content type
                page_width = Inches(8.5)  # Standard page width minus margins
                total_width = Inches(7.5)  # Usable width
              
                # Calculate column widths - optimized for better distribution with smaller S. No.
                col_widths = {}
                for col_name in df.columns:
                    if col_name == 'S. No.':
                        col_widths[col_name] = 288000  # 0.2 inches in EMUs (reduced from 0.3)
                    elif col_name == 'Vessel Name':
                        col_widths[col_name] = 1728000  # 1.2 inches in EMUs
                    elif col_name == 'Comments':
                        col_widths[col_name] = 5760000  # 4.0 inches in EMUs
                    elif col_name == 'Potential Fuel Saving (MT/Day)':
                        col_widths[col_name] = 1152000  # 0.8 inches in EMUs
                    elif col_name == 'YTD CII':
                        col_widths[col_name] = 576000  # 0.4 inches in EMUs
                    elif 'Hull Condition' in col_name and 'Power Loss' not in col_name:
                        col_widths[col_name] = 864000  # 0.6 inches in EMUs
                    elif 'Excess Power %' in col_name:
                        col_widths[col_name] = 864000  # 0.6 inches in EMUs
                    elif 'ME Efficiency' in col_name:
                        col_widths[col_name] = 864000  # 0.6 inches in EMUs
                    else:
                        col_widths[col_name] = 864000  # 0.6 inches in EMUs
              
                # Set column widths using integer values
                for i, col_name in enumerate(df.columns):
                    table.columns[i].width = col_widths[col_name]
              
                # Style header row
                header_cells = table.rows[0].cells
                for i, col_name in enumerate(df.columns):
                    cell = header_cells[i]
                    cell.text = col_name
                  
                    # Header styling
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)
                  
                    # Header background color
                    set_cell_shading(cell, "2F75B5")
                  
                    # Header borders
                    set_cell_border(
                        cell,
                        top={"sz": 6, "val": "single", "color": "000000"},
                        left={"sz": 6, "val": "single", "color": "000000"},
                        bottom={"sz": 6, "val": "single", "color": "000000"},
                        right={"sz": 6, "val": "single", "color": "000000"}
                    )
              
                # Add data rows
                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        cell = row_cells[i]
                        cell_value = str(value) if pd.notna(value) else "N/A"
                        cell.text = cell_value
                      
                        # Cell styling based on column type
                        column_name = df.columns[i]
                      
                        # Set font size
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                      
                        # Apply conditional formatting
                        if 'Hull Condition' in column_name and 'Power Loss' not in column_name:
                            # Hull Condition text columns (Good/Average/Poor)
                            color_hex = get_cell_color(cell_value)
                            if color_hex:
                                set_cell_shading(cell, color_hex)
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif 'Excess Power %' in column_name:
                            # Power Loss percentage columns (numeric values)
                            try:
                                if pd.notna(value):
                                    power_loss_num = float(value)
                                    # Format as number with 1 decimal place
                                    cell.text = f"{power_loss_num:.1f}"
                                    # Apply color based on numeric value
                                    if power_loss_num < 15:
                                        set_cell_shading(cell, "D4EDDA")  # Green
                                    elif 15 <= power_loss_num <= 25:
                                        set_cell_shading(cell, "FFF3CD")  # Yellow
                                    else:
                                        set_cell_shading(cell, "F8D7DA")  # Red
                            except (ValueError, TypeError):
                                pass
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif 'ME Efficiency' in column_name:
                            color_hex = get_cell_color(cell_value)
                            if color_hex:
                                set_cell_shading(cell, color_hex)
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif column_name == 'YTD CII':
                            # Apply CII background color coding with black text
                            cii_bg_color = get_cii_background_color(cell_value)
                            if cii_bg_color:
                                set_cell_shading(cell, cii_bg_color)
                            # Set text to black and bold
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
                                    run.font.bold = True
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif column_name == 'Comments':
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                      
                        # Set text wrapping and vertical alignment
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                      
                        # Enable text wrapping
                        tcW = OxmlElement("w:tcW")
                        tcW.set(qn("w:type"), "auto")
                        tcPr.append(tcW)
                      
                        # Set vertical alignment to top
                        vAlign = OxmlElement('w:vAlign')
                        vAlign.set(qn('w:val'), 'top')
                        tcPr.append(vAlign)
                      
                        # Add borders
                        set_cell_border(
                            cell,
                            top={"sz": 6, "val": "single", "color": "000000"},
                            left={"sz": 6, "val": "single", "color": "000000"},
                            bottom={"sz": 6, "val": "single", "color": "000000"},
                            right={"sz": 6, "val": "single", "color": "000000"}
                        )
              
                # Add page break before appendix
                doc.add_page_break()
              
                # Add Appendix section
                appendix_title = doc.add_paragraph()
                appendix_run = appendix_title.add_run("Appendix")
                appendix_run.font.size = Pt(20)
                appendix_run.font.bold = True
                appendix_run.font.color.rgb = RGBColor(255, 255, 255)
                appendix_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
              
                # Set appendix title background
                appendix_title_format = appendix_title.paragraph_format
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:val'), 'clear')
                shading_elm.set(qn('w:color'), 'auto')
                shading_elm.set(qn('w:fill'), '00B0F0')
                appendix_title_format._element.get_or_add_pPr().append(shading_elm)
              
                # Add spacing
                doc.add_paragraph("")
              
                # General Conditions section
                general_heading = doc.add_paragraph()
                general_run = general_heading.add_run("General Conditions")
                general_run.font.size = Pt(14)
                general_run.font.bold = True
                general_run.font.color.rgb = RGBColor(47, 117, 181)
              
                # Add bullet points for general conditions
                conditions = [
                    "Analysis Period is Last Six Months or after the Last Event whichever is later",
                    "Days with Good Weather (BF<=4) are considered for analysis",
                    "Days with Steaming hrs greater than 17 considered for analysis",
                    "Data is compared with Original Sea Trial"
                ]
              
                for condition in conditions:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.first_line_indent = Inches(-0.25)
                    run = p.add_run("• " + condition)
                    run.font.size = Pt(10)
              
                # Hull Performance section
                doc.add_paragraph("")
                hull_heading = doc.add_paragraph()
                hull_run = hull_heading.add_run("Hull Performance")
                hull_run.font.size = Pt(14)
                hull_run.font.bold = True
                hull_run.font.color.rgb = RGBColor(47, 117, 181)
              
                # Hull performance criteria with colors
                hull_criteria = [
                    ("Excess Power < 15% – Rating Good", RGBColor(0, 176, 80)),
                    ("15% < Excess Power < 25% – Rating Average", RGBColor(255, 192, 0)),
                    ("Excess Power > 25% – Rating Poor", RGBColor(255, 0, 0))
                ]
              
                for criteria, color in hull_criteria:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.first_line_indent = Inches(-0.25)
                    run = p.add_run("• " + criteria)
                    run.font.size = Pt(10)
                    run.font.color.rgb = color
              
                # Machinery Performance section
                doc.add_paragraph("")
                machinery_heading = doc.add_paragraph()
                machinery_run = machinery_heading.add_run("Machinery Performance")
                machinery_run.font.size = Pt(14)
                machinery_run.font.bold = True
                machinery_run.font.color.rgb = RGBColor(47, 117, 181)
              
                # Machinery performance criteria with colors
                machinery_criteria = [
                    ("SFOC (g/kWh) within ±10 from Shop test condition are considered as \"Good\"", RGBColor(0, 176, 80)),
                    ("SFOC (g/kWh) Greater than 10 and less than 20 are considered as \"Average\"", RGBColor(255, 192, 0)),
                    ("SFOC (g/kWh) Above 20 are considered as \"Poor\"", RGBColor(255, 0, 0))
                ]
              
                for criteria, color in machinery_criteria:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.first_line_indent = Inches(-0.25)
                    run = p.add_run("• " + criteria)
                    run.font.size = Pt(10)
                    run.font.color.rgb = color

                # CII Rating section
                # doc.add_paragraph("")
                # cii_heading = doc.add_paragraph()
                # cii_run = cii_heading.add_run("CII Rating")
                # cii_run.font.size = Pt(14)
                # cii_run.font.bold = True
                # cii_run.font.color.rgb = RGBColor(47, 117, 181)
              
                # CII performance criteria with colors
                # cii_criteria = [
                #     ("Rating A – Significantly Better Performance", RGBColor(144, 238, 144)),  # Light green
                #     ("Rating B – Better Performance", RGBColor(0, 100, 0)),                    # Dark green
                #     ("Rating C – Moderate Performance", RGBColor(255, 215, 0)),               # Yellow
                #     ("Rating D – Minor Inferior Performance", RGBColor(255, 140, 0)),         # Orange
                #     ("Rating E – Inferior Performance", RGBColor(255, 0, 0))                  # Red
                # ]
              
                # for criteria, color in cii_criteria:
                #     p = doc.add_paragraph()
                #     p.paragraph_format.left_indent = Inches(0.25)
                #     p.paragraph_format.first_line_indent = Inches(-0.25)
                #     run = p.add_run("• " + criteria)
                #     run.font.size = Pt(10)
                #     run.font.color.rgb = color
              
                break
      
        if not placeholder_found:
            st.warning("Placeholder '{{Template}}' not found. Adding report at the end of document.")
            # If no placeholder found, add content at the end
            doc.add_page_break()
          
            # Add the same content as above but at the end
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run("Fleet Performance Analysis Report")
            title_run.font.size = Pt(24)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(47, 117, 181)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
          
            # Continue with the rest of the report...
      
        # Save to bytes buffer
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
      
    except Exception as e:
        st.error(f"Error creating Word report: {str(e)}")
        return None

# Function to reset session state
def reset_page():
    st.session_state.selected_vessels = set()
    st.session_state.report_data = None
    st.session_state.search_query = ""
    st.session_state.report_months = 2
    st.cache_data.clear()
    st.session_state.trigger_rerun = True

# Main Application
def main():
    # Check for rerun flag
    if 'trigger_rerun' in st.session_state and st.session_state.trigger_rerun:
        st.session_state.trigger_rerun = False
        st.rerun()

    # Lambda Function URL
    LAMBDA_FUNCTION_URL = "https://yrgj6p4lt5sgv6endohhedmnmq0eftti.lambda-url.ap-south-1.on.aws/"

    # Enhanced Header with gradient styling
    st.markdown("""
    <div class="main-header">
        <h1>🚢 Enhanced Vessel Performance Report Tool</h1>
        <p>Generate comprehensive performance reports with advanced analytics and beautiful formatting</p>
    </div>
    """, unsafe_allow_html=True)

    # Reset button with enhanced styling
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("🔄 Reset All", type="secondary"):
            reset_page()

    # Load vessels with office mapping
    with st.spinner("Loading vessels..."):
        try:
            vessel_directory = fetch_vessel_directory(LAMBDA_FUNCTION_URL)
            all_vessels = [record["vessel_name"] for record in vessel_directory]
            vessel_office_lookup = {record["vessel_name"]: record["office"] for record in vessel_directory}
            office_to_vessels = defaultdict(list)
            for record in vessel_directory:
                office_to_vessels[record["office"]].append(record["vessel_name"])
            office_options = ["All Offices"] + sorted(office_to_vessels.keys())
            st.markdown(f'<div class="success-box">✅ Successfully loaded {len(all_vessels)} vessels!</div>', unsafe_allow_html=True)
        except Exception as e:
            st.error(f"❌ Failed to load vessels: {str(e)}")
            vessel_directory = []
            all_vessels = []
            vessel_office_lookup = {}
            office_to_vessels = defaultdict(list)
            office_options = ["All Offices"]

    summary_placeholder = st.container()
    selection_tab, report_tab, analytics_tab, guide_tab = st.tabs(
        ["🎯 DOC & Selection", "🚀 Reports", "📊 Analytics", "📖 Guide"]
    )

    if not all_vessels:
        with selection_tab:
            st.error("❌ Failed to load vessels. Please check your connection and try again.")
        with report_tab:
            st.info("Load vessels in the DOC tab to generate a report.")
        with analytics_tab:
            st.info("Analytics will appear once a report is generated.")
        with guide_tab:
            st.info("Refer to the documentation below.")
        return

    # --- Selection Tab ---
    with selection_tab:
        st.markdown('<div class="section-header">🎯 DOC & Vessel Selection</div>', unsafe_allow_html=True)

        if st.session_state.selected_office not in office_options:
            st.session_state.selected_office = "All Offices"

        office_col1, office_col2 = st.columns([3, 1])
        with office_col1:
            selected_office = st.selectbox(
                "🏢 Choose DOC",
                office_options,
                index=office_options.index(st.session_state.selected_office),
                help="Selecting a DOC auto-selects every vessel under it.",
            )
        with office_col2:
            office_count = len(office_to_vessels.get(selected_office, all_vessels)) if selected_office != "All Offices" else len(all_vessels)
            st.metric("DOC Vessels", office_count)

        if selected_office != st.session_state.selected_office:
            st.session_state.selected_office = selected_office
            if selected_office != "All Offices":
                st.session_state.selected_vessels = set(office_to_vessels.get(selected_office, []))
            st.session_state.checkbox_version += 1

        search_query = st.text_input(
            "🔍 Search vessels",
            value=st.session_state.search_query,
            placeholder="Type vessel name to filter...",
            help="Start typing to filter the vessel list in real-time"
        )

        if search_query != st.session_state.search_query:
            st.session_state.search_query = search_query

        base_pool = office_to_vessels.get(selected_office, all_vessels) if selected_office != "All Offices" else all_vessels
        filtered_vessels = filter_vessels_client_side(base_pool, search_query)

        metric_col1, metric_col2, metric_col3 = st.columns(3)
        with metric_col1:
            st.markdown(f"""
            <div class="metric-card">
                <h3>📊 Total Vessels</h3>
                <h2>{len(all_vessels)}</h2>
            </div>
            """, unsafe_allow_html=True)
        with metric_col2:
            st.markdown(f"""
            <div class="metric-card">
                <h3>🔍 Filtered</h3>
                <h2>{len(filtered_vessels)}</h2>
            </div>
            """, unsafe_allow_html=True)
        with metric_col3:
            st.markdown(f"""
            <div class="metric-card">
                <h3>✅ Selected</h3>
                <h2>{len(st.session_state.selected_vessels)}</h2>
            </div>
            """, unsafe_allow_html=True)

        if filtered_vessels:
            st.subheader("Select Vessels")
            with st.container(height=320, border=True):
                cols = st.columns(3)
                for i, vessel in enumerate(filtered_vessels):
                    col_idx = i % 3
                    checkbox_state = cols[col_idx].checkbox(
                        f"{vessel} ({vessel_office_lookup.get(vessel, 'Unassigned DOC')})",
                        value=(vessel in st.session_state.selected_vessels),
                        key=f"checkbox_{st.session_state.checkbox_version}_{vessel}"
                    )
                    if checkbox_state:
                        st.session_state.selected_vessels.add(vessel)
                    else:
                        if vessel in st.session_state.selected_vessels:
                            st.session_state.selected_vessels.remove(vessel)
        else:
            st.markdown('<div class="info-box">🔍 No vessels match your search criteria</div>', unsafe_allow_html=True)

        selected_vessels_list = list(st.session_state.selected_vessels)
        if selected_vessels_list:
            with st.expander(f"📋 Selected Vessels ({len(selected_vessels_list)})", expanded=False):
                for i, vessel in enumerate(sorted(selected_vessels_list), 1):
                    st.write(f"{i}. {vessel}")

        with st.expander("🩺 Diagnostics", expanded=False):
            if not st.session_state.diagnostics:
                st.caption("No diagnostics recorded yet.")
            else:
                for entry in st.session_state.diagnostics:
                    status_emoji = "✅" if entry["status"] == "success" else "⚠️"
                    st.write(f"{status_emoji} [{entry['time']}] **{entry['label']}** — {entry['detail']}")

        st.session_state.summary_metrics = {
            "doc": selected_office,
            "total": len(all_vessels),
            "filtered": len(filtered_vessels),
            "selected": len(st.session_state.selected_vessels),
        }
        render_summary_bar(summary_placeholder)

    selected_vessels_list = list(st.session_state.selected_vessels)

    # Enhanced report duration selection
    st.subheader("📅 Select Report Duration:")
    st.session_state.report_months = st.radio(
        "Report Duration",
        options=[1, 2, 3],
        format_func=lambda x: f"📊 {x} Month{'s' if x > 1 else ''} Analysis",
        index=1,
        horizontal=True,
        help="Choose the number of months for historical analysis",
        label_visibility="collapsed"
    )

    if selected_vessels_list:
        # Enhanced generate button
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("🚀 Generate Performance Report", type="primary"):
                with st.spinner("Generating comprehensive report with advanced analytics..."):
                    try:
                        start_time = time.time()
                        st.session_state.report_data = query_report_data(
                            LAMBDA_FUNCTION_URL, selected_vessels_list, st.session_state.report_months
                        )

                        generation_time = time.time() - start_time

                        if not st.session_state.report_data.empty:
                            st.markdown(f'<div class="success-box">✅ Report generated successfully in {generation_time:.2f} seconds!</div>', unsafe_allow_html=True)
                        else:
                            st.warning("⚠️ No data found for the selected vessels.")

                    except Exception as e:
                        st.error(f"❌ Error generating report: {str(e)}")
                        st.session_state.report_data = None
    else:
        st.markdown('<div class="info-box">⚠️ Please select at least one vessel to generate a report</div>', unsafe_allow_html=True)

    # Section 3: Report Display and Download
    if st.session_state.report_data is not None and not st.session_state.report_data.empty:
        st.markdown('<div class="section-header">3. 📊 Report Results & Analytics</div>', unsafe_allow_html=True)

        # Enhanced report display
        st.subheader("📋 Performance Data Table")
        styled_df = st.session_state.report_data.style.apply(style_condition_columns, axis=1)
        st.dataframe(styled_df, height=400, width="stretch")

        # Enhanced download section
        st.subheader("📥 Download Options")
        col1, col2, col3 = st.columns(3)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        with col1:
            filename = f"vessel_performance_report_{timestamp}.xlsx"
            try:
                excel_data = create_excel_download_with_styling(st.session_state.report_data, filename)
                if excel_data:
                    st.download_button(
                        label="📊 Download Excel Report",
                        data=excel_data,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    )
            except Exception as e:
                st.error(f"❌ Error creating Excel file: {str(e)}")

        with col2:
            csv_data = st.session_state.report_data.to_csv(index=False)
            csv_filename = f"vessel_performance_report_{timestamp}.csv"
            st.download_button(
                label="📄 Download CSV Report",
                data=csv_data,
                file_name=csv_filename,
                mime="text/csv",
            )

        with col3:
            word_filename = f"fleet_performance_report_{timestamp}.docx"
            try:
                template_path = "Fleet Performance Template.docx"
                if os.path.exists(template_path):
                    word_data = create_enhanced_word_report(st.session_state.report_data, template_path, st.session_state.report_months)
                    if word_data:
                        st.download_button(
                            label="📝 Download Word Report",
                            data=word_data,
                            file_name=word_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    else:
                        st.error("❌ Failed to create Word report")
                else:
                    st.warning("⚠️ Template file not found")
            except Exception as e:
                st.error(f"❌ Error creating Word file: {str(e)}")

        # Enhanced analytics section
        with st.expander("📈 Advanced Analytics & Insights", expanded=False):
            tab1, tab2, tab3, tab4 = st.tabs(["🛡️ Hull Analysis", "⚙️ Engine Analysis", "📊 Trend Analysis", "🌍 CII Analysis"])

            with tab1:
                st.subheader("Hull Condition Distribution")
                hull_cols = [col for col in st.session_state.report_data.columns if 'Hull Condition' in col]

                if hull_cols:
                    col1, col2 = st.columns(2)
                    with col1:
                        latest_hull_data = st.session_state.report_data[hull_cols[0]].value_counts()
                        if len(latest_hull_data) > 0:
                            st.bar_chart(latest_hull_data, width="stretch")
                            st.caption(f"Distribution for {hull_cols[0]}")

                    with col2:
                        hull_summary = []
                        for col in hull_cols:
                            month = col.replace("Hull Condition ", "")
                            counts = st.session_state.report_data[col].value_counts()
                            hull_summary.append({
                                "Month": month,
                                "Good": counts.get("Good", 0),
                                "Average": counts.get("Average", 0),
                                "Poor": counts.get("Poor", 0),
                                "N/A": counts.get("N/A", 0)
                            })
                      
                        if hull_summary:
                            hull_summary_df = pd.DataFrame(hull_summary)
                            st.dataframe(hull_summary_df, width="stretch")

            with tab2:
                st.subheader("ME Efficiency Distribution")
                me_cols = [col for col in st.session_state.report_data.columns if 'ME Efficiency' in col]

                if me_cols:
                    col1, col2 = st.columns(2)
                    with col1:
                        latest_me_data = st.session_state.report_data[me_cols[0]].value_counts()
                        if len(latest_me_data) > 0:
                            st.bar_chart(latest_me_data, width="stretch")
                            st.caption(f"Distribution for {me_cols[0]}")

                    with col2:
                        me_summary = []
                        for col in me_cols:
                            month = col.replace("ME Efficiency ", "")
                            counts = st.session_state.report_data[col].value_counts()
                            me_summary.append({
                                "Month": month,
                                "Good": counts.get("Good", 0),
                                "Average": counts.get("Average", 0),
                                "Poor": counts.get("Poor", 0),
                                "Anomalous": counts.get("Anomalous data", 0),
                                "N/A": counts.get("N/A", 0)
                            })
                      
                        if me_summary:
                            me_summary_df = pd.DataFrame(me_summary)
                            st.dataframe(me_summary_df, width="stretch")

            with tab3:
                st.subheader("Performance Trends")
              
                hull_cols = [col for col in st.session_state.report_data.columns if 'Hull Condition' in col]
                me_cols = [col for col in st.session_state.report_data.columns if 'ME Efficiency' in col]

                if len(hull_cols) >= 2:
                    st.write("**Hull Condition Trends (% Good)**")
                    hull_trend_data = []
                    for col in hull_cols:
                        month = col.replace("Hull Condition ", "")
                        total_with_data = len(st.session_state.report_data[st.session_state.report_data[col] != "N/A"])
                        good_count = len(st.session_state.report_data[st.session_state.report_data[col] == "Good"])
                      
                        if total_with_data > 0:
                            percentage = (good_count / total_with_data * 100)
                            hull_trend_data.append({"Month": month, "Good %": percentage})
                  
                    if hull_trend_data:
                        hull_trend_df = pd.DataFrame(hull_trend_data)
                        if hull_trend_df["Good %"].sum() > 0:
                            st.line_chart(hull_trend_df.set_index("Month"), width="stretch")

                if len(me_cols) >= 2:
                    st.write("**ME Efficiency Trends (% Good)**")
                    me_trend_data = []
                    for col in me_cols:
                        month = col.replace("ME Efficiency ", "")
                        total_with_data = len(st.session_state.report_data[st.session_state.report_data[col] != "N/A"])
                        good_count = len(st.session_state.report_data[st.session_state.report_data[col] == "Good"])
                      
                        if total_with_data > 0:
                            percentage = (good_count / total_with_data * 100)
                            me_trend_data.append({"Month": month, "Good %": percentage})
                  
                    if me_trend_data:
                        me_trend_df = pd.DataFrame(me_trend_data)
                        if me_trend_df["Good %"].sum() > 0:
                            st.line_chart(me_trend_df.set_index("Month"), width="stretch")

            with tab4:
                st.subheader("CII Rating Distribution")
                if 'YTD CII' in st.session_state.report_data.columns:
                    col1, col2 = st.columns(2)
                    with col1:
                        cii_data = st.session_state.report_data['YTD CII'].value_counts()
                        if len(cii_data) > 0:
                            st.bar_chart(cii_data, width="stretch")
                            st.caption("CII Rating Distribution")

                    with col2:
                        cii_summary = []
                        counts = st.session_state.report_data['YTD CII'].value_counts()
                        total_vessels = len(st.session_state.report_data)
                        
                        for rating in ['A', 'B', 'C', 'D', 'E']:
                            count = counts.get(rating, 0)
                            percentage = (count / total_vessels * 100) if total_vessels > 0 else 0
                            cii_summary.append({
                                "Rating": rating,
                                "Count": count,
                                "Percentage": f"{percentage:.1f}%"
                            })
                        
                        # Add N/A if exists
                        na_count = counts.get('N/A', 0) + sum(1 for x in st.session_state.report_data['YTD CII'] if pd.isna(x))
                        if na_count > 0:
                            na_percentage = (na_count / total_vessels * 100) if total_vessels > 0 else 0
                            cii_summary.append({
                                "Rating": "N/A",
                                "Count": na_count,
                                "Percentage": f"{na_percentage:.1f}%"
                            })
                        
                        if cii_summary:
                            cii_summary_df = pd.DataFrame(cii_summary)
                            st.dataframe(cii_summary_df, width="stretch")
                else:
                    st.info("No CII data available for analysis")

    # Enhanced Help Section
    with st.expander("📖 User Guide & Features", expanded=False):
        st.markdown("""
        ### 🌟 Enhanced Features:

        **🎨 Modern UI Design:**
        - Gradient backgrounds and modern styling
        - Responsive layout with improved visual hierarchy
        - Color-coded metrics and status indicators

        **🔍 Smart Vessel Selection:**
        - Real-time search and filtering
        - Multi-column layout for easy browsing
        - Visual metrics showing selection status

        **📊 Advanced Analytics:**
        - Interactive charts and visualizations
        - Multi-month trend analysis
        - Performance distribution insights
        - CII rating analysis with color coding

        **📥 Professional Reports:**
        - Enhanced Excel reports with color coding
        - Beautifully formatted Word documents
        - Optimized table layouts with proper spacing
        - CII rating color coding in all formats

        ### 📋 How to Use:

        1. **🔍 Search**: Use the search box to find specific vessels
        2. **✅ Select**: Check vessels you want to analyze
        3. **📅 Configure**: Choose analysis period (1-3 months)
        4. **🚀 Generate**: Click to create comprehensive report
        5. **📊 Analyze**: Review charts and performance metrics
        6. **📥 Download**: Export in your preferred format

        ### 🎯 Performance Indicators:

        **🛡️ Hull Condition:**
        - 🟢 **Good**: < 15% excess power
        - 🟡 **Average**: 15-25% excess power
        - 🔴 **Poor**: > 25% excess power

        **⚙️ Engine Efficiency:**
        - 🟢 **Good**: 160-180 g/kWh SFOC
        - 🟡 **Average**: 180-190 g/kWh SFOC
        - 🔴 **Poor**: > 190 g/kWh SFOC
        - ⚪ **Anomalous**: < 160 g/kWh SFOC
        """)

if __name__ == "__main__":
    main()
